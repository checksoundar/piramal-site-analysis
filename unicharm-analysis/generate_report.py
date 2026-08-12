#!/usr/bin/env python3
"""Generate the Unicharm Global Web Estate — AEM to Edge Delivery Services
migration analysis Word document (RFP deliverable)."""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
EMBED = os.path.join(BASE, "embed")
SITES = json.load(open(os.path.join(BASE, "sites.json")))["sites"]

BRAND = RGBColor(0x00, 0x6C, 0xB5)      # Unicharm blue
BRAND_HEX = "006CB5"
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x70, 0x70, 0x70)
LIGHT_ROW = "EAF3FA"

CAT_NAME = {
    "CP": "Corporate", "WC": "Wellness Care", "PC": "Pet Care",
    "FC": "Feminine Care", "BC": "Baby Care",
}


# ------------------------------------------------------------------ helpers
def shade(cell, color):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), color)
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def add_table(doc, headers, rows, widths=None, font=8.0, header_font=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(header_font)
        r.font.color.rgb = RGBColor(255, 255, 255)
        shade(hdr[i], BRAND_HEX)
    set_repeat_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run("" if val is None else str(val))
            r.font.size = Pt(font)
            if ri % 2 == 1:
                shade(cells[ci], LIGHT_ROW)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def body(doc, text, size=10, space=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(size)


def add_shot(doc, sid, caption, width=6.0, kind="full"):
    fp = os.path.join(EMBED, f"{sid}_{kind}.jpg")
    if not os.path.exists(fp):
        fp = os.path.join(EMBED, f"{sid}_top.jpg")
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        return True
    return False


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.color.rgb = BRAND if level <= 2 else DARK
    return hd


# ================================================================== build
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

# ------- cover
for _ in range(5):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Unicharm Global Web Estate")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = BRAND
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("AEM → Adobe Edge Delivery Services\nMigration Analysis & Assessment")
r.font.size = Pt(17)
r.font.color.rgb = DARK
doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(
    "68 confirmed-AEM migration-target sites  ·  22 countries / regions  ·  5 brand families\n"
    "Prepared in response to the Unicharm RFP URL list (2026-08-03)")
r.font.size = Pt(11)
r.font.color.rgb = GREY
doc.add_paragraph()
d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("Migration Assessment Report  ·  August 2026")
r.font.size = Pt(11)
r.font.color.rgb = GREY
doc.add_page_break()

# ------- TOC
h(doc, "Table of Contents", 1)
for item in [
    "1. Executive Summary",
    "2. Templates Inventory",
    "3. Blocks / Components Catalog",
    "4. Page Counts by Template",
    "5. Integrations Analysis",
    "6. Complex Use Cases & Observations",
    "7. Migration Estimates",
    "8. Appendix A: Site Inventory (68 sites)",
    "9. Appendix B: Per-Site Screenshots",
]:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)
doc.add_page_break()

# ================================================================== 1. EXEC
h(doc, "1. Executive Summary", 1)
body(doc,
     "This document assesses the migration of Unicharm's global web estate to Adobe Edge Delivery "
     "Services (EDS / aem.live). Scope, per the RFP URL list dated 2026-08-03, is the 68 properties "
     "confirmed to run on Adobe Experience Manager (AEM) and flagged as migration targets "
     "(AEMか = 〇, 移行対象か = 〇). These span 22 countries/regions and five brand "
     "families — Corporate (Unicharm), Baby Care (MamyPoko, Moony, Bobby, BabyJoy, Torepanman, "
     "Oyasumiman), Wellness Care (Lifree, Caryn, Carenavi), Feminine Care (Sofy, Charm, Diana, "
     "Silcot, Charmnap, Sofygirls) and Pet Care (Unicharm Pet).")
body(doc,
     "All sampled sites were verified as traditional server-rendered AEM (author/publish). Evidence: "
     "Digital Asset Manager paths (/content/dam/sites/{site}/), client-library references "
     "(/etc.clientlibs/), jcr namespaces, and a systematic numbered component library. Critically, the "
     "estate is delivered by two shared AEM platforms built by two vendors, and each platform reuses "
     "one component set across all its regional sites — which is the single largest driver of migration "
     "efficiency.")

h(doc, "Platform Split", 2)
add_table(doc,
          ["Platform", "Brand Families", "Sites", "Known Pages", "Component Library Markers"],
          [
              ["Vendor A", "Corporate (CP), Baby Care (BC), Pet Care (PC)", "39",
               "6,562", "uc-mod-* / uc-js-* modules, CMP-CM## numbered components"],
              ["Vendor B", "Wellness Care (WC), Feminine Care (FC)", "29",
               "3,785", "CMP-CM## / HC## / FC## / GL## components, slick.js carousel"],
              ["TOTAL", "5 brand families", "68", "~10,350 (60 sites) + est.", "Shared per-vendor libraries"],
          ],
          widths=[2.6, 6.2, 1.6, 2.4, 5.0])

h(doc, "Key Findings", 2)
bullets(doc, [
    "68 in-scope AEM sites across 22 countries; ~10,350 pages confirmed across 60 sites (8 sites without a published count), for an estimated ~10,700-page total portfolio.",
    "Two shared component libraries (one per vendor) span corporate + all brand + RTL sites — template and block reuse is very high. The correct estimating unit is BLOCKS, not sites.",
    "~13 distinct page templates per platform, assembled from a common palette of ~25 reusable blocks dominated by 3–4 content primitives (heading / text / image-and-text / media card).",
    "Content is highly standardized and largely bulk-migratable: article/advice/news/product-detail pages follow uniform, numbered URL patterns ideal for automated import (est. ~65% automatic).",
    "Effort concentrates in 6 large sites holding ~6,245 pages (≈60% of known content): Unicharm Pet JP (1,428), MamyPoko Indonesia (1,400), Unicharm Corporate JP (1,317), Moony JP (900), Sofy JP (700), Lifree JP (500).",
    "10 right-to-left (RTL) Arabic sites (Unicharm SA/EG, Sofyclub SA/EG, BabyJoy SA/EG/AE/BH/KW/OM) — a cross-cutting requirement affecting every block, larger than initially assumed.",
    "Single global Google Tag Manager container (GTM-WLCGVS) across all sites; Adobe Experience Platform Launch on higher-traffic JP/SEA brands. Analytics migrates cheaply via delayed loading.",
    "No third-party consent-management platform (CMP) was found — the 'cmp-' strings are AEM Core Components, not a consent tool. CMP compliance across 22 jurisdictions is a NEW workstream, not a port.",
    "Custom interactive tools require bespoke re-build: multi-retailer 'buy now' EC modals (region-specific retailer sets), menstruation cycle calendar (Diana), diaper size-finder (Moony), birthday/age personalization (MamyPoko/BabyJoy), and Sofygirls 'like' widgets (implies a backend state service).",
    "Portfolio-wide jQuery + slick.js dependency must be removed and reimplemented as vanilla-JS EDS blocks to hit the PageSpeed-100 / Core Web Vitals target.",
    "8 sites are marked for consolidation/decommission into 3 targets (Mask + Wave → unicharm.co.jp; Torepanman + Oyasumiman → jp.moony.com; Gulf BabyJoy AE/BH/KW/OM → arabia.babyjoyclub.com).",
    "One site (BabyJoy Egypt, www.babyjoy.com.eg) is currently unreachable due to a server-side TLS/SNI misconfiguration on its Imperva CDN — flagged as an availability/access risk for migration.",
])
doc.add_page_break()

# ================================================================== 2. TEMPLATES
h(doc, "2. Templates Inventory", 1)
body(doc,
     "Both platforms compose pages from a shared component palette rather than rigid per-type page "
     "types, so 'templates' below are the recurring page archetypes observed across the estate. "
     "Templates are largely common between the two vendors; where a vendor-specific pattern exists it "
     "is noted. Complexity reflects the effort to reproduce the template in Edge Delivery.")

h(doc, "2.1 Vendor A Templates (Corporate, Baby Care, Pet Care)", 2)
add_table(doc,
          ["Template", "Complexity", "Reasoning", "Example URL(s)"],
          [
              ["TA-01 Homepage / Brand Top", "Complex",
               "Hero carousel + many teaser/card rows + news feeds + brand switcher; per-site variants.",
               "www.unicharm.co.jp/\njp.moony.com/ja/home.html"],
              ["TA-02 Product Listing / Category", "Medium",
               "Grid of product cards + category nav + EC-purchase entry points.",
               "jp.moony.com/ja/products.html\nid.mamypoko.com/id/products.html"],
              ["TA-03 Product Detail", "Complex",
               "Image gallery, size chart, weight-range tables, per-retailer buy modals, related products, video.",
               "jp.moony.com/ja/products/mnm.html"],
              ["TA-04 News / Press Listing", "Medium",
               "Year-filterable dated news list (CMP-CM34 list module).",
               "www.unicharm.co.jp/ja/company/news.html"],
              ["TA-05 News / Article Detail", "Simple–Medium",
               "Heading + dated body + breadcrumb + SNS share; light component count.",
               "www.unicharm.co.jp/ja/company/news/2026/0206-01.html"],
              ["TA-06 Editorial / Tips Article", "Medium",
               "Long-form childcare articles: image+text modules, accordions, related-article cards.",
               "jp.moony.com/ja/tips/baby/childcare/diapers/bm0025.html"],
              ["TA-07 Corporate Content (About/CSR/IR)", "Medium",
               "Structured heading/text/table/figure modules; deep IR & CSR trees, PDF libraries.",
               "www.unicharm.co.jp/ja/ir/library.html\n/ja/csr-eco.html"],
              ["TA-08 Campaign / Landing", "Medium",
               "Banner + promo cards + CTA buttons; marketing-driven, lighter structure.",
               "jp.moony.com/ja/campaign.html"],
              ["TA-09 Where-to-Buy / Online-Shop Hub", "Complex",
               "Product picker + size buttons + expandable per-retailer buy menus + EC modals; heavy repetition.",
               "id.mamypoko.com/id/shop.html"],
              ["TA-10 Size Guide / Diaper Selector", "Medium",
               "Weight/size mapping tables + product recommendation modules.",
               "jp.moony.com/ja/diapers.html"],
              ["TA-11 Product / Tips Search Results", "Medium",
               "Form inputs + facets + product/article result cards.",
               "jp.moony.com/ja/products/search.html"],
              ["TA-12 FAQ / Contact", "Medium",
               "Contact tables + accordions; some FAQ offloaded to legacy domains.",
               "www.unicharm.co.jp/ja/contact-us.html"],
              ["TA-13 Corporate Site Search", "Complex",
               "MarsFlag hosted search API (corporate only).",
               "www.unicharm.co.jp/ (header search)"],
          ],
          widths=[3.6, 1.8, 6.6, 5.4])

h(doc, "2.2 Vendor B Templates (Wellness Care, Feminine Care)", 2)
add_table(doc,
          ["Template", "Complexity", "Reasoning", "Example URL(s)"],
          [
              ["TB-01 Homepage", "Complex",
               "Hero carousel + news headline + multiple promo grids + brand list.",
               "jp.lifree.com/ja/home.html\nwww.sofy.jp/ja/home.html"],
              ["TB-02 Product Listing / Category", "Medium",
               "Repeated product-list block (HC02/FC05) + heading/link-list.",
               "jp.lifree.com/ja/product/list.html\nwww.sofy.jp/ja/products.html"],
              ["TB-03 Category by Attribute (ADL level)", "Medium",
               "Listing filtered by ADL 0–5 (Lifree adult-incontinence-specific).",
               "jp.lifree.com/ja/product/adl1.html"],
              ["TB-04 Product Detail", "Complex",
               "Rich text + video + EC buy-button + product registration + spec grids.",
               "www.sofy.jp/ja/products/organic/organic_napkin.html"],
              ["TB-05 Article / Advice Detail", "Medium",
               "Heavy image-and-text + heading sequence + related-content registration.",
               "www.sofy.jp/ja/advice/during-period/01.html\nwww.carenavi.jp/ja/basic/omutsu/choice/reduce.html"],
              ["TB-06 Article / Advice Listing (Hub)", "Medium",
               "Card grid (grid-advance) + link lists.",
               "jp.lifree.com/ja/advice.html\nwww.carenavi.jp/ja/basic.html"],
              ["TB-07 Editorial Column Detail", "Medium",
               "Article layout + multi-link module (Carenavi expert/genba/report series).",
               "www.carenavi.jp/ja/expert/vol15.html"],
              ["TB-08 Column / Volume Index", "Simple",
               "Container + grid-advance list of volumes.",
               "www.carenavi.jp/ja/expert/list.html"],
              ["TB-09 Product Finder / Selection Guide", "Medium",
               "Composed choice cards linking to filtered listings (navigation, not a JS quiz).",
               "jp.lifree.com/ja/product/adult/choose.html\nwww.sofy.jp/ja/products/product-type.html"],
              ["TB-10 FAQ", "Medium",
               "Accordion via grid-advance (toggle state); 100+ Q&A text nodes.",
               "jp.lifree.com/ja/advice/qa.html"],
              ["TB-11 Campaign / Landing", "Medium",
               "Standalone promo pages under /campaign/.",
               "www.sofy.jp/ja/campaign/pinkribbon.html"],
              ["TB-12 Menstruation Cycle Tool", "Complex",
               "Interactive calendar/cycle process page with client-side date math (Diana).",
               "www.diana.com.vn/vi/calendar/process.html"],
              ["TB-13 Contact / Sitemap / Static", "Simple",
               "Text + link-list only.",
               "www.carenavi.jp/ja/contact.html"],
          ],
          widths=[3.6, 1.8, 6.6, 5.4])

body(doc,
     "Note: neither platform exposes an on-site global search-results template on the brand sites "
     "(product buying delegates to external e-commerce; corporate uses the hosted MarsFlag search). "
     "'Where to buy' on Vendor B sites links to the proprietary map.unicharm.co.jp locator and retailer "
     "search rather than an embedded map.", size=9, italic=True, color=GREY)
doc.add_page_break()

# ================================================================== 3. BLOCKS
h(doc, "3. Blocks / Components Catalog", 1)
body(doc,
     "The catalog groups visual variations of the same content model as ONE block with variants noted, "
     "per the analysis brief. Component IDs in the description are the actual AEM markers observed in "
     "the delivered markup. The two platforms share most content primitives; brand-specific header/"
     "footer variants are grouped.")

h(doc, "3.1 Shared / Core Blocks (both platforms)", 2)
add_table(doc,
          ["Block", "Complexity", "Behaviour & Functionality", "Where Used"],
          [
              ["Global Header / Mega-Menu", "Complex",
               "Dropdown mega-menu (Company/Sustainability/IR/Recruit/Products on corporate; brand nav on brand sites) + language & 'Worldwide Site' region selector + mobile hamburger. Per-brand variants (uc-js-header-nav; CM02/HC04/FC01/GL).",
               "All 68 sites"],
              ["Global Footer", "Medium",
               "Sitemap link columns + brand list + policy/social links; collapsible accordion on mobile (uc-mny-js-footer-accordion; CM04/HC05/FC02).",
               "All 68 sites"],
              ["Hero / Main-Visual Carousel", "Complex",
               "Auto-rotating slides, prev/next + play/pause; responsive desktop/mobile images. Variants: corporate carousel-00, brand carousel-02, slick-based CM13 slide-banner.",
               "All homepages; product & campaign pages"],
              ["Teaser / Media Card", "Medium",
               "Image+text/media card — the workhorse block (highest count on nearly every page). Used for banners, related content, category tiles (CMP-CM10/CM11, uc-mod-media-01).",
               "Every template, all sites"],
              ["Heading / Text / Rich-Text", "Simple",
               "Core default-content primitives — heading levels, paragraphs, formatted body (CM05/CM09/CM10/CM34).",
               "Every page"],
              ["Layout Grid / Section Container", "Simple",
               "Structural grid/row/column wrappers (parsys) (uc-lyt-*, CM05-block/CM06-section/CM32/CM33).",
               "Every page"],
              ["Breadcrumb", "Simple",
               "Hierarchical path navigation (CM19; uc-mod-nav-breadcrumb-01).",
               "All deep pages"],
              ["Accordion / Toggle", "Medium",
               "Expand/collapse for FAQ, product specs, mobile footer; jQuery toggle with open state (uc-js-tgl-01; CM40 grid-advance).",
               "FAQ, product, contact"],
              ["Video Embed", "Medium",
               "Embedded YouTube/inline player (CM18-video; youtube.com/embed).",
               "Product detail, tips, campaign"],
              ["EC 'Buy Here' Purchase Modal", "Complex",
               "Expandable per-retailer buy menu aggregating multiple e-commerce retailers per product/size; region-specific retailer arrays (uc-mod-modal-ec, uc-js-modal-inline; CM26-ec-button).",
               "Product detail, listing, shop hub (baby & feminine)"],
              ["Product Card / Grid", "Medium",
               "Product tile: image, name, size badges (uc-mod-product-01/02; HC02/FC05 product-list).",
               "Listing, search, home"],
              ["Product / Content Registration", "Medium",
               "Registers product/article context to drive dynamic related-content listings (HC01/FC07; FC09).",
               "Product & advice detail (Vendor B)"],
              ["News / Article List", "Medium",
               "Dated, year-filterable list (CMP-CM34; CM24 news headline).",
               "Corporate news/CSR/IR, homepages"],
              ["Social Share / Follow", "Simple",
               "Share and follow icons; AddToAny and native SNS (CM16/CM31; uc-mod-sns-01; addtoany).",
               "Most pages"],
              ["Language / Region Selector", "Simple",
               "Locale switcher → /global.html country gateway (CM20).",
               "All sites"],
              ["CTA Button", "Simple",
               "Styled call-to-action button (uc-mod-btn-01; CM12).",
               "Everywhere"],
              ["Reference / Experience Fragment", "Simple",
               "Includes shared content (header/footer reuse) (CM28-reference).",
               "All pages"],
              ["Raw HTML injection", "Simple",
               "Author-injected custom markup (CM15-html) — migration wildcard, needs per-instance review.",
               "All sites"],
          ],
          widths=[3.4, 1.7, 8.0, 4.2])

h(doc, "3.2 Specialized / Interactive Blocks", 2)
add_table(doc,
          ["Block", "Complexity", "Behaviour & Functionality", "Where Used"],
          [
              ["Online-Shop / Retailer Directory", "Complex",
               "Product picker → size → retailer picker with sticky nav; up to 47 EC-modal instances on one page (uc-mod-shop-online-product-01, uc-mod-sticky-root).",
               "MamyPoko shop hubs"],
              ["Diaper Size Chart / Finder", "Medium",
               "Weight-to-size mapping tables (kg ranges → S/M/L/Big) + product recommendation (uc-mod-size-chart-01).",
               "Moony, MamyPoko diaper pages"],
              ["Birthday / Child-Age Personalization", "Medium",
               "Month/day/year selector driving age-relevant content (uc-mod-set-birthday-wrapper, uc-poko-list).",
               "MamyPoko, BabyJoy"],
              ["Menstruation Cycle Calendar", "Complex",
               "Form with cycle length/period inputs + client-side date calculation (fields mensDaysPeriod, mensDaysLength).",
               "Diana (VN); Sofy SofyBe app integration"],
              ["Product Finder / Selection Cards", "Medium",
               "Composed choice cards routing to filtered listings; also Lifree ADL-level selector.",
               "Sofy, Lifree"],
              ["'Like' / Engagement Widget", "Complex",
               "Client-side like button + persisted count (implies backend/state service) (uc-js-like-btn/-cnt; CM44).",
               "Sofygirls (youth site)"],
              ["Modal / Lightbox", "Medium",
               "Popup dialogs and image lightboxes (CM14; uc-js-modal-01, uc-js-media-square-01).",
               "Sofygirls, EC modals, media grids"],
              ["App-Store Badge", "Simple",
               "iOS / Google Play download links (CM37).",
               "Sofygirls, Sofy (SofyBe)"],
              ["Cookie Consent Banner (basic)", "Simple",
               "Lightweight consent banner (CM38-general-cookie) on selected sites; NOT a full CMP.",
               "Lifree (limited)"],
              ["Back-to-Top", "Simple",
               "Scroll-to-top control (utl-page-top).",
               "Various"],
          ],
          widths=[3.4, 1.7, 8.0, 4.2])
doc.add_page_break()

# Illustrative screenshots for templates/blocks
h(doc, "3.3 Illustrative Screenshots", 2)
body(doc, "Representative homepages illustrating the shared template and block system. "
          "A complete per-site screenshot set is provided in Appendix B.", size=9, italic=True, color=GREY)
add_shot(doc, "s01", "Figure 1: Unicharm Corporate JP (Vendor A) — TA-01 Homepage, mega-menu, hero carousel, triple news feed.", 6.2, "top")
add_shot(doc, "s04", "Figure 2: Moony JP (Vendor A, Baby Care) — hero carousel, product cards, campaign & tips blocks.", 6.2, "top")
doc.add_page_break()
add_shot(doc, "s08", "Figure 3: Lifree JP (Vendor B, Wellness Care) — TB-01 Homepage, ADL selection guide, care advice.", 6.2, "top")
add_shot(doc, "s13", "Figure 4: Sofy JP (Vendor B, Feminine Care) — hero, product lineup, editorial advice blocks.", 6.2, "top")
doc.add_page_break()
add_shot(doc, "s55", "Figure 5: BabyJoy Saudi Arabia (Vendor A) — RTL Arabic layout reusing the same component set.", 6.2, "top")
add_shot(doc, "s54", "Figure 6: Sofyclub Saudi Arabia (Vendor B) — RTL Arabic feminine-care site.", 6.2, "top")
doc.add_page_break()

# ================================================================== 4. PAGE COUNTS
h(doc, "4. Page Counts by Template", 1)
body(doc,
     "Page counts derive from the RFP's published per-domain totals (ページ枚数 sheet). Because the "
     "estate is composed from a shared block palette, counts are best reasoned per template archetype "
     "using the portfolio-wide content mix rather than per individual site. The mix below reflects the "
     "observation that article/advice/news/product-detail pages (standardized, numbered URLs) dominate "
     "the estate and are automatable, while homepages, hubs, tools and search pages are manual.")

h(doc, "4.1 Portfolio Content Mix by Template Class", 2)
add_table(doc,
          ["Template Class", "Est. Pages", "% of Estate", "Migration Type", "Rationale"],
          [
              ["Article / Advice / Editorial detail (TA-06, TB-05/07)", "~4,200", "~39%", "Automatic",
               "Uniform heading/text/image-text composition; numbered URLs; ideal bulk import."],
              ["News / Press detail & listing (TA-04/05)", "~1,400", "~13%", "Automatic",
               "Standardized dated articles; year-filtered indexes regenerated from data."],
              ["Product detail (TA-03, TB-04)", "~1,500", "~14%", "Semi-Automatic",
               "Structured data imports, but EC buy-modals, size charts & video need review."],
              ["Product listing / category / finder (TA-02, TB-02/03/09)", "~900", "~8%", "Semi-Automatic",
               "Template + card grids; finder/ADL cards are navigation, low risk."],
              ["Corporate content — About/CSR/IR (TA-07)", "~700", "~7%", "Semi-Automatic",
               "Rich structured content + PDF libraries; IR/CSR governance overhead."],
              ["Campaign / landing (TA-08, TB-11)", "~500", "~5%", "Semi-Automatic",
               "Marketing pages; bespoke layouts but block-composable."],
              ["FAQ / Contact / Static / Sitemap (TA-12, TB-10/13)", "~600", "~6%", "Semi-Automatic",
               "Accordions + link lists; mostly mechanical."],
              ["Homepages / Brand tops (TA-01, TB-01)", "~70", "~1%", "Manual",
               "Complex multi-carousel assembly; ~1 per site + sub-brand tops."],
              ["Where-to-buy / Shop hubs (TA-09)", "~120", "~1%", "Manual",
               "Heavy retailer-directory logic; region-specific EC modals."],
              ["Interactive tools & search (TA-10/11/13, TB-12)", "~150", "~1%", "Manual",
               "Size finder, cycle calendar, MarsFlag search, like widgets — custom JS."],
              ["Unclassified / long-tail buffer", "~530", "~5%", "Mixed",
               "Miscellaneous & 8 sites without published counts (est. @40 pp)."],
              ["TOTAL (estimated portfolio)", "~10,700", "100%", "—",
               "~10,350 confirmed across 60 sites + estimates for 8 uncounted sites."],
          ],
          widths=[5.4, 1.9, 1.6, 2.3, 6.0])

h(doc, "4.2 Migration Automatability Summary", 2)
add_table(doc,
          ["Migration Category", "Est. Pages", "% of Estate", "Includes"],
          [
              ["Automatic (bulk import)", "~6,900", "~65%",
               "Article/advice/editorial, news/press. Standardized, numbered URLs — script-driven import."],
              ["Semi-Automatic (template + review)", "~2,700", "~25%",
               "Product detail/listing, corporate/CSR/IR, campaign, FAQ/static. Imported then manually validated."],
              ["Manual (custom development)", "~1,100", "~10%",
               "Homepages, shop hubs, interactive tools, search, RTL edge cases, consolidation targets."],
          ],
          widths=[4.4, 2.0, 1.8, 9.0])
body(doc, "Automatability is high overall because ~52% of the estate is article and news content "
          "with a near-uniform three-primitive structure. The manual tail is small in page count but "
          "concentrated in the highest-complexity blocks.", size=9, italic=True, color=GREY)

h(doc, "4.3 Largest Sites (effort concentration)", 2)
big = sorted([s for s in SITES if s["pages"]], key=lambda x: -x["pages"])[:10]
add_table(doc,
          ["Site", "Brand Family", "Vendor", "Pages", "Dominant Migration Type"],
          [[s["brand"] + " (" + s["country"] + ")", CAT_NAME[s["cat"]], s["vendor"], str(s["pages"]),
            "Automatic (content-heavy)" if s["pages"] >= 300 else "Mixed"] for s in big],
          widths=[4.2, 3.0, 1.4, 1.6, 5.0])
body(doc, "The six sites above 500 pages hold ~6,245 pages — roughly 60% of all known content. "
          "Migration scheduling and QA cost are dominated by these properties; the ~20 sites under 65 "
          "pages are low-effort or consolidation candidates.", size=9, italic=True, color=GREY)
doc.add_page_break()

# ================================================================== 5. INTEGRATIONS
h(doc, "5. Integrations Analysis", 1)
body(doc,
     "Integrations below are backed by literal strings found in the delivered HTML of representative "
     "sites (raw fetches across JP corporate + brands, SEA, Arabic, India, and AU). Where a common "
     "integration was NOT found, that is stated explicitly, as it affects the migration scope.")
add_table(doc,
          ["Integration", "Type", "Complexity", "Evidence (ID / domain)", "Where Found"],
          [
              ["Google Tag Manager", "Analytics / Tag mgmt", "Simple",
               "GTM-WLCGVS (single global container on all sites)", "All 68 sites"],
              ["Adobe Experience Platform Launch", "Analytics / Tag mgmt", "Medium",
               "assets.adobedtm.com/launch-EN… (distinct embed per brand)",
               "Moony, Sofy JP, Lifree, MamyPoko ID/PH"],
              ["GA4 / Google Ads / pixels", "Analytics / Ad-pixel", "Simple",
               "No hard-coded G-/UA-/AW- IDs — fired at runtime via GTM/Launch", "(runtime, via GTM)"],
              ["Facebook / Meta SDK", "Social", "Medium",
               "connect.facebook.net/…/sdk.js appId=270586063403176", "www.unicharm.co.jp"],
              ["YouTube embeds (IFrame API)", "Embed", "Simple",
               "youtube.com/embed/…, youtube.com/iframe_api", "unicharm.co.jp, sofy.jp, sa.sofyclub.com"],
              ["MarsFlag site search", "API / Search", "Complex",
               "ce.mf.marsflag.com + finder.api.mf.marsflag.com/…/search", "www.unicharm.co.jp (corporate)"],
              ["Store locator (proprietary)", "Maps", "Complex",
               "map.unicharm.co.jp/?brand=Bxxxx&sub_brand=SBxxxx", "unicharm.co.jp, moony, sofy, lifree"],
              ["Shindan diagnosis chatbot", "Chat", "Complex",
               "shindan1.c.unicharm.com/bot_switcher/… (session-tokenized)", "Lifree, corporate refs"],
              ["Moony point / member program", "API / Member", "Complex",
               "point.moony.com, app.moony.com/… (stateful, separate host)", "jp.moony.com"],
              ["E-commerce 'where to buy' links (JP)", "E-commerce", "Simple",
               "amazon.co.jp, lohaco.yahoo.co.jp, rakuten.co.jp (plain links)", "sofy.jp, lifree, moony"],
              ["E-commerce links (SEA)", "E-commerce", "Simple",
               "shopee.co.id, lazada.com.ph, tokopedia, blibli (plain links)", "id.mamypoko.com, ph.mamypoko.com"],
              ["AddToAny social share", "Social", "Simple",
               "static.addtoany.com", "unicharm.co.jp, lifree, .com.au, .co.in"],
              ["jQuery 3.6.0 + touchSwipe", "Library", "Simple (remove)",
               "ajax.googleapis.com/…/jquery-3.6.0 + jquery.touchSwipe", "All sites"],
              ["Slick carousel", "Library", "Simple (remove)",
               "cdn.jsdelivr.net/…/slick-carousel@1.8.1", "Vendor B (Lifree et al.)"],
              ["Consent Management Platform (CMP)", "CMP", "Complex (NEW)",
               "NOT FOUND — no OneTrust/Optanon/Cookiebot; 'cmp-' = AEM Core Components", "None (compliance gap)"],
              ["reCAPTCHA", "API", "—",
               "NOT FOUND in sampled homepage HTML", "None sampled"],
          ],
          widths=[3.9, 2.4, 1.9, 5.6, 4.0])

h(doc, "5.1 Martech Stack Summary", 2)
bullets(doc, [
    "Google + Adobe dual-tagged, centrally governed: one global GTM container (GTM-WLCGVS) on every site worldwide injects GA4/Ads/pixels dynamically (hence no hard-coded analytics IDs in HTML).",
    "Adobe Experience Platform Launch adds a second layer on higher-traffic brands (Moony, Sofy JP, Lifree, MamyPoko ID/PH), each with its own launch-EN… embed — implies selective Adobe Analytics/Target use.",
    "Regional gradient: Japan is the richest stack (GTM + Launch + MarsFlag search + map.unicharm.co.jp locator + Moony point program + shindan chatbot + Facebook SDK); SEA adds marketplace links + Adobe Launch; Arabic sites are minimal (GTM + social + occasional YouTube); India & Australia are the thinnest (GTM + AddToAny), functioning largely as brand-directory hubs.",
    "Front-end libraries are uniform: jQuery 3.6.0 (Google-hosted) + touchSwipe everywhere; slick.js on Vendor B. No Google Fonts/Typekit detected (self-hosted fonts).",
])

h(doc, "5.2 Migration Notes", 2)
body(doc, "Trivial to carry over (load via delayed.js or author as links/blocks):", size=10)
bullets(doc, [
    "GTM (single container covers all locales; GA4/Ads/pixels ride inside it — migrate for free).",
    "Adobe Launch (one async script per brand, preserving each distinct launch-EN… embed).",
    "AddToAny, social-profile links, e-commerce 'where to buy' links, YouTube embeds.",
])
body(doc, "Must be dropped and reimplemented in vanilla JS for the EDS performance target:", size=10)
bullets(doc, ["jQuery + touchSwipe + slick.js → native EDS carousel / interaction blocks."])
body(doc, "Needs real re-engineering or a dedicated embed/loader block:", size=10)
bullets(doc, [
    "MarsFlag search API → new EDS search block wired to the endpoint (or replacement).",
    "map.unicharm.co.jp store locator → re-embed or rebuild as a locator block.",
    "Shindan chatbot → dedicated embed with back-end coordination.",
    "Moony point/member program → out of scope as an app, but entry-point CTAs/links must be preserved.",
    "Facebook SDK / XFBML widgets → replace with lazy social embeds.",
    "CMP: no consent platform currently exists — adding one (e.g. OneTrust) across 22 jurisdictions is a NEW compliance workstream, not a port.",
])
doc.add_page_break()

# ================================================================== 6. COMPLEX
h(doc, "6. Complex Use Cases & Observations", 1)
body(doc, "Counts are marked [C] confirmed against the inventory or live pages, or [E] estimated / "
          "extrapolated. They identify functionality needing special attention during migration.")
add_table(doc,
          ["Use Case", "Instances", "Where Found", "Why Complex", "Migration Impact / Recommendation"],
          [
              ["Multi-retailer 'buy now' EC modal aggregation", "~40+ commerce sites [E]; 47 modals on one page [C]",
               "id.mamypoko.com/id/shop.html; mamypoko.jp home; product pages (baby & feminine)",
               "Per-product modal opens a retailer set that differs by region (JP vs SEA marketplaces); each SKU carries its own retailer array.",
               "Build ONE EDS buy-now block driven by a per-locale retailer JSON/spreadsheet. Highest-reuse block in the program."],
              ["Interactive stateful JS tools", "5+ tool types across ~15+ sites [E]",
               "Diana cycle calendar [C]; Moony diaper size-finder [C]; Lifree ADL selector; MamyPoko/BabyJoy birthday widget; Sofy fertility sheets",
               "Application logic (form state, date math, conditional recommendations), currently jQuery-driven and server-page-backed.",
               "Each needs a bespoke vanilla-JS EDS block; no auto-import path. Scope as individual custom line items."],
              ["RTL Arabic layout sites", "10 sites [C]",
               "Unicharm SA/EG, Sofyclub SA/EG, BabyJoy SA/EG/AE/BH/KW/OM (html dir=rtl verified)",
               "Full mirrored layout + bidirectional text; carousels, modals & nav must be RTL-aware. EDS global CSS is LTR-first.",
               "Establish an RTL CSS baseline (logical properties, dir-scoped rules) once; test every block in RTL. Cross-cutting requirement."],
              ["Multi-region / multi-language + brand-name localization", "22 countries, ~15 languages [C]; 5 brand mappings [C]",
               "Sofy=Charm=Diana; MamyPoko=Bobby=BabyJoy; Lifree=Caryn",
               "Same product platform, divergent brand identity, domain, language and retailer set; per-locale paths + per-site DAM trees.",
               "Shared block library + per-brand theme (logo, tokens, fonts) + per-locale content. Do NOT fork code per brand."],
              ["Site consolidation / decommission mergers", "8 source sites → 3 targets [C]",
               "Mask+Wave → unicharm.co.jp; Torepanman+Oyasumiman → jp.moony.com; BabyJoy AE/BH/KW/OM → arabia.babyjoyclub.com",
               "URL restructuring, 301 redirect maps, DAM merges & IA reconciliation on top of migration; Gulf merge folds 4 RTL sites into 1.",
               "Deliver redirect maps + consolidated IA per target; add a country/store selector for the Gulf merge. Migrate target first, then fold in sources."],
              ["Large-volume bulk-migration sites", "6 sites = 6,245 pages (~60%) [C]",
               "Unicharm Pet 1,428; MamyPoko ID 1,400; Unicharm JP 1,317; Moony 900; Sofy 700; Lifree 500",
               "Need automated import pipelines, template detection & product-catalog handling at scale; QA cost dominated by these.",
               "Invest in import automation + template inventory for the top 6; ~60% of migration effort. Tiny sites (2–15 pp) hand-rebuilt or absorbed."],
              ["Corporate special sections (IR / CSR / News / Recruit)", "unicharm.co.jp + 13 corporate sites [C]",
               "unicharm.co.jp /ir/ (calendar, library, e-announcement), /csr-eco/ (ESG/GRI/SASB data), year-based /company/news/ archive",
               "Financial-document / PDF libraries, structured ESG datasets, multi-year date-filtered archives, regulatory sensitivity.",
               "Needs a document-library block, a filterable news index with year facets, and an ESG data-table pattern. Highest governance/accessibility burden."],
              ["jQuery + slick carousel dependency", "Every sampled site [C]",
               "All platforms (shared numbered uc-js-* / CM## libraries)",
               "EDS targets zero-dependency, LCP-optimized vanilla JS; jQuery + slick block the PageSpeed-100 target and won't import cleanly.",
               "Reimplement carousels as one vanilla-JS, RTL-aware, lazy EDS block; eliminate jQuery entirely. Portfolio-wide dependency removal."],
              ["Youth interactive site (Sofygirls)", "1 site, 340 pp [C]",
               "jp.sofygirls.com — like buttons + counts, 59 modal refs, app-store badges",
               "Client-side 'like' counters imply a backend/state service beyond static content; richest interactive surface.",
               "Requires a like/engagement micro-service or 3rd-party integration (EDS is static-first). Confirm counts are real vs cosmetic before estimating."],
              ["Legacy-domain offloads (interfaced, out of scope)", "Referenced from in-scope sites [C]",
               "map.unicharm.co.jp, shindan bot, faq.unicharm.co.jp, unicharmgame.com",
               "Not migrating, but migrated pages link INTO them — broken-link / mixed-UX risk if not inventoried.",
               "Inventory all outbound links to legacy domains; preserve as external links. Mark explicitly out-of-scope in the SOW."],
              ["Forms + consent / CMP compliance", "Forms on tool/contact pages [C]; no CMP tool [C]",
               "Diana calendar form; GTM on all pages; NO OneTrust/Optanon/Cookiebot found",
               "22 jurisdictions (EU-adjacent, Gulf, SEA, Brazil LGPD) but consent appears GTM-managed only; forms need backend endpoints EDS lacks natively.",
               "Verify per-region consent obligations; likely add a CMP during migration. Route forms to a forms service / AEM Forms. Treat consent as a NEW workstream."],
              ["BabyJoy Egypt unreachable (TLS/SNI)", "1 site [C]",
               "www.babyjoy.com.eg — ERR_SSL_UNRECOGNIZED_NAME_ALERT (Imperva CDN)",
               "Server-side SNI misconfiguration; site cannot be crawled/captured by any client until fixed.",
               "Flag as availability risk; obtain content via origin/CMS export. Template is shared with other BabyJoy sites, so design is not blocked."],
          ],
          widths=[3.1, 2.2, 3.6, 4.0, 4.6], font=7.5, header_font=8.0)
doc.add_page_break()

# ================================================================== 7. ESTIMATES
h(doc, "7. Migration Estimates", 1)
body(doc,
     "Estimates assume the block-first strategy that the shared component libraries make possible: build "
     "a reusable EDS block set once, then run automated content import across the estate. A blended team "
     "of EDS developers, a content-migration engineer, QA, and a project lead is assumed. Ranges reflect "
     "the low/high of scope decisions (e.g. how many interactive tools are rebuilt vs retired).")

h(doc, "7.1 Effort by Workstream", 2)
add_table(doc,
          ["Workstream", "Scope", "Effort (person-days)"],
          [
              ["Discovery & design system", "Design-token extraction per brand family (5), RTL baseline, template/block architecture, redirect strategy.", "25–35"],
              ["Core block library", "~18 shared blocks (header/mega-menu, footer, hero carousel, media card, product card, accordion, news index, breadcrumb, social, language selector, etc.), RTL-aware, vanilla JS.", "60–85"],
              ["Specialized / interactive blocks", "EC buy-now modal, shop directory, size-finder, cycle calendar, birthday widget, product finder, like widget, MarsFlag search, store-locator embed, document library.", "55–80"],
              ["Automated content import", "Import pipeline + template detection for ~6,900 auto pages (articles/news), incl. the 6 large sites; parsers/transformers per template.", "45–65"],
              ["Semi-automatic migration & review", "~2,700 pages (product, corporate/CSR/IR, campaign, FAQ) — import then manual validation & fix-up.", "50–70"],
              ["Manual page builds", "~1,100 pages: homepages, shop hubs, tool pages, search, RTL edge cases, consolidation targets.", "45–65"],
              ["Site consolidation / redirects", "8 source → 3 target merges: IA reconciliation, 301 maps, DAM merges, Gulf country selector.", "15–22"],
              ["Integrations & martech", "GTM/Launch re-wiring, social/YouTube embeds, EC links, MarsFlag/locator/chatbot embeds, member-program CTAs.", "18–26"],
              ["CMP & compliance (new)", "Select & implement a consent-management platform across 22 jurisdictions; forms routing.", "15–25"],
              ["QA & testing", "Cross-browser, responsive, RTL, bilingual/CJK+Arabic, performance (PageSpeed-100), link & redirect validation.", "55–80"],
              ["Project management & UAT support", "Coordination across 5 brand families / 2 vendor stacks, stakeholder reviews, launch waves.", "35–50"],
              ["Contingency (~15%)", "Risk buffer (BabyJoy EG access, tool scope, backend services).", "55–75"],
          ],
          widths=[4.4, 9.6, 3.2])

h(doc, "7.2 Totals", 2)
add_table(doc,
          ["Metric", "Low", "High"],
          [
              ["Total effort (person-days)", "473", "678"],
              ["Total effort (person-months, ~21 pd/mo)", "~22.5", "~32.3"],
              ["Calendar schedule (blended 5–7 person team, phased waves)", "~7 months", "~10 months"],
              ["Indicative cost @ blended US$550/day*", "~US$260,000", "~US$373,000"],
          ],
          widths=[8.0, 4.6, 4.6])
body(doc, "*Indicative day-rate for planning only; actual pricing depends on delivery location, team mix "
          "and contract terms. Cost excludes AEM/EDS licensing, third-party CMP/search subscriptions, "
          "and translation services (content is already localized).", size=8.5, italic=True, color=GREY)

h(doc, "7.3 Recommended Phasing (waves)", 2)
add_table(doc,
          ["Wave", "Focus", "Scope & Outcome"],
          [
              ["Wave 1 (Mo 1–2)", "Foundation",
               "Design system + RTL baseline, core block library, import pipeline. Pilot: one large JP brand (e.g. Sofy or Moony)."],
              ["Wave 2 (Mo 2–5)", "Bulk content + large sites",
               "Automated import of the 6 large sites (~6,245 pp) and article/news content. Semi-automatic product/corporate pages."],
              ["Wave 3 (Mo 4–7)", "Brand + regional rollout",
               "Remaining SEA/global brand sites; specialized blocks (EC modal, finders, calendar); integrations & CMP."],
              ["Wave 4 (Mo 6–9)", "RTL, consolidation & tools",
               "10 RTL Arabic sites, 8→3 consolidation with redirects, interactive tools & search, backend services."],
              ["Wave 5 (Mo 8–10)", "QA, UAT & launch",
               "Full-estate performance/RTL/bilingual QA, redirect validation, phased go-live."],
          ],
          widths=[3.2, 3.2, 10.8])
doc.add_page_break()

# ================================================================== 8. APPENDIX A
h(doc, "8. Appendix A: Site Inventory (68 sites)", 1)
body(doc, "All confirmed-AEM migration-target sites from the RFP URL list, with vendor platform "
          "assignment, brand family, published page count, and notes. '—' = no page count published.",
     size=9, italic=True, color=GREY)
inv_rows = []
for i, s in enumerate(SITES, 1):
    inv_rows.append([
        str(i), s["country"], CAT_NAME[s["cat"]], s["brand"],
        s["url"].replace("https://", "").replace("http://", ""),
        s["vendor"], str(s["pages"]) if s["pages"] is not None else "—",
    ])
add_table(doc,
          ["#", "Country/Region", "Family", "Brand", "URL", "Vendor", "Pages"],
          inv_rows,
          widths=[0.8, 2.6, 2.2, 2.8, 6.4, 1.2, 1.2], font=7.5, header_font=8.0)
doc.add_page_break()

# ================================================================== 9. APPENDIX B
h(doc, "9. Appendix B: Per-Site Screenshots", 1)
body(doc, "Homepage screenshot for each in-scope site (captured August 2026, 1440px viewport). "
          "One site (BabyJoy Egypt) could not be captured due to a server-side TLS/SNI error and is "
          "noted below. Screenshots are grouped by vendor platform then country.",
     size=9, italic=True, color=GREY)

# Order: Vendor A then B, by country order as in inventory
ordered = [s for s in SITES if s["vendor"] == "A"] + [s for s in SITES if s["vendor"] == "B"]
fig = 7
missing = []
for s in ordered:
    cap = (f"Figure {fig}: {s['brand']} — {s['country']} · {CAT_NAME[s['cat']]} · Vendor {s['vendor']} · "
           f"{s['url'].replace('https://','').replace('http://','')}"
           + (f" · {s['pages']} pp" if s['pages'] else ""))
    ok = add_shot(doc, s["id"], cap, 5.6, "top")
    if ok:
        fig += 1
        doc.add_paragraph()
    else:
        missing.append(s)
if missing:
    body(doc, "Not captured: " + "; ".join(f"{m['brand']} ({m['url']})" for m in missing)
         + " — server-side TLS/SNI misconfiguration (Imperva CDN); content available via CMS export.",
         size=9, italic=True, color=GREY)

# ---- footer note on final page
doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run("Prepared from live inspection of representative pages across both vendor platforms "
                 "(August 2026). Page counts sourced from the client RFP. This assessment is a planning "
                 "estimate; a fixed bid requires a detailed template/block workshop and CMS content export.")
r.italic = True
r.font.size = Pt(8)
r.font.color.rgb = GREY

out = os.path.join(BASE, "Unicharm_Global_Web_Estate_Migration_Analysis.docx")
doc.save(out)
print("Saved:", out)
print("Size: %.1f KB" % (os.path.getsize(out) / 1024))
print("Figures embedded:", fig - 7, "of", len(ordered), "sites; missing:", [m["id"] for m in missing])
