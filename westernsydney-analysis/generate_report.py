#!/usr/bin/env python3
"""Western Sydney University — AEM → Edge Delivery Services migration analysis (RFP deliverable)."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
EMBED = os.path.join(BASE, "embed")

BRAND = RGBColor(0xA0, 0x00, 0x30)   # WSU crimson
BRAND_HEX = "A00030"
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x70, 0x70, 0x70)
LIGHT = "F5E9EE"


def shade(cell, color):
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), color); sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)


def add_table(doc, headers, rows, widths=None, font=8.5, hfont=9):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"; t.autofit = False
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(hfont); r.font.color.rgb = RGBColor(255, 255, 255)
        shade(c, BRAND_HEX)
    repeat_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""; r = cells[ci].paragraphs[0].add_run("" if val is None else str(val)); r.font.size = Pt(font)
            if ri % 2 == 1: shade(cells[ci], LIGHT)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    return t


def body(doc, text, size=10, space=6, italic=False, color=None, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); r.font.size = Pt(size)


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = BRAND if level <= 2 else DARK
    return hd


def shot(doc, key, caption, width=6.2, kind="top"):
    fp = os.path.join(EMBED, f"{key}_{kind}.jpg")
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(width)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY


doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10)

# ---------- cover ----------
for _ in range(5): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Western Sydney University"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = BRAND
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Website Migration Analysis\nAEM → Adobe Edge Delivery Services"); r.font.size = Pt(17); r.font.color.rgb = DARK
doc.add_paragraph()
su = doc.add_paragraph(); su.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = su.add_run("https://www.westernsydney.edu.au/\n9,925 pages analysed  ·  August 2026"); r.font.size = Pt(11); r.font.color.rgb = GREY
doc.add_page_break()

# ---------- TOC ----------
h(doc, "Contents", 1)
for item in ["1. Executive Summary", "2. Templates Inventory", "3. Blocks / Components Catalog",
             "4. Page Counts by Template", "5. Integrations Analysis", "6. Complex Use Cases & Observations",
             "Appendix: Template Screenshots"]:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(4)
    for run in p.runs: run.font.size = Pt(11)
doc.add_page_break()

# ---------- 1. Exec ----------
h(doc, "1. Executive Summary", 1)
body(doc,
     "This document analyses the Western Sydney University (WSU) website for migration to Adobe Edge "
     "Delivery Services (EDS / aem.live). Analysis is based on the site's XML sitemap (9,925 pages), "
     "programmatic URL-pattern analysis, live inspection of representative pages across every template "
     "family, and rendered-DOM component extraction.")
body(doc, "Headline finding — the migration is already substantially underway.", bold=True, color=BRAND, space=3)
body(doc,
     "WSU runs two platforms in parallel behind one domain, and the direction of travel is clear from the "
     "evidence. The deep content estate — research-institute and centre microsites, schools, the entire "
     "News Centre, events and people pages (the ~9,900 sitemap URLs) — is ALREADY on Edge Delivery: pages "
     "declare a meta tag cms_type=EDS, load the EDS boilerplate (/scripts/aem.js, /scripts/scripts.js), "
     "serve EDS block tables via the .plain.html endpoint, and emit EDS RUM telemetry (rum.hlx.page). A "
     "random sample of 40 sitemap URLs returned 40/40 EDS. The residual classic-AEM surface is the "
     "corporate marketing shell — the homepage and top-level pages such as /future and /research "
     "(cms_type=AEM, ~109 /etc.clientlibs/ references, cmp-* Core Components), the corporate navigation "
     "and footer experience fragments, the bespoke news-events widget, the Algolia search integration, "
     "and AEM Adaptive Forms.")
h(doc, "Key Findings", 2)
bullets(doc, [
    "Platform: Adobe Experience Manager, mid-transition to Edge Delivery Services. Confirmed via clientlibs, cmp-* components, EDS boilerplate scripts, .plain.html block tables, and cms_type meta tags.",
    "9,925 pages in the sitemap — effectively all already rendered by Edge Delivery. The classic-AEM corporate pages are largely NOT in the sitemap and form a separate, much smaller set.",
    "~11 template archetypes, dominated by research-institute/centre microsites (40+ centres sharing a common sub-page pattern: about-us, our-people, research, projects, publications, events, news).",
    "~1,870 News-story detail pages and ~884 event pages — highly standardised, ideal for automated handling.",
    "~15 reusable blocks. The EDS target vocabulary (hero, cards, carousel, columns, table, accordions, dynamic news-list, nav/footer fragments) already exists and is proven in production across thousands of pages — this de-risks the remaining work.",
    "Navigation and footer on EDS pages are authored as content fragments (per-microsite nav pages + a shared /footer), rather than hard-coded — a clean, maintainable pattern.",
    "Extensive martech: a single Adobe DTM/Launch library orchestrates two GTM containers, Adobe Audience Manager, and an unusually large ad-pixel farm (Meta ×5 IDs, TikTok ×4, Snapchat, LinkedIn, Google Ads ×3, Floodlight, Spotify) aimed at student recruitment. Consent via OneTrust.",
    "Higher-effort integrations to re-engineer: Algolia search (5 production indices incl. courses), AEM Adaptive Forms, Genesys PureCloud live chat, and Qualtrics intercepts.",
    "Enterprise systems on separate subdomains (handbook, apply, library→LibCal, login→Okta, goglobal, careers) are out of CMS scope; the site only links to them.",
    "Access note: the site returns HTTP 406 to non-browser agents; all analysis used browser headers.",
])
doc.add_page_break()

# ---------- 2. Templates ----------
h(doc, "2. Templates Inventory", 1)
body(doc, "Template archetypes observed across the estate. 'Stack' notes whether the template is already on "
          "Edge Delivery (EDS) or remains on classic AEM. Complexity reflects effort to (re)build/verify the "
          "template in the EDS target state.")
add_table(doc,
          ["Template", "Stack", "Complexity", "Reasoning", "Example URL(s)"],
          [
              ["T01 Corporate Homepage", "AEM", "Complex",
               "Only page still on classic AEM Core Components; bespoke news-events widget (news+events+calendar+gallery), promo bands, embedded Adaptive Forms, Algolia.",
               "https://www.westernsydney.edu.au/"],
              ["T02 Corporate Marketing Page", "AEM", "Complex",
               "Top-level recruitment pages on classic AEM; rich hero/promo layouts + Algolia course search + pixels.",
               "https://www.westernsydney.edu.au/future\nhttps://www.westernsydney.edu.au/research"],
              ["T03 Microsite Landing", "EDS", "Medium",
               "Institute/centre/school landing: per-microsite nav fragment + hero + carousel + cards + dynamic news list. Repeats across 40+ centres.",
               "https://www.westernsydney.edu.au/ics\nhttps://www.westernsydney.edu.au/schools/grs"],
              ["T04 Generic Content / Rich-Text", "EDS", "Simple",
               "Default content (headings, paragraphs, links) + occasional cards/columns/table. The long tail (about-us, policy, scholarships, info pages).",
               "https://www.westernsydney.edu.au/ics/about-us"],
              ["T05 News Centre Hub", "EDS", "Medium",
               "Combines dynamic news-list (query/limit/pagination) with curated table-based story lists.",
               "https://www.westernsydney.edu.au/news-centre"],
              ["T06 News Story Detail", "EDS", "Simple-Medium",
               "Rich-text body + inline images + optional accordions. ~1,870 pages, highly uniform.",
               "https://www.westernsydney.edu.au/news-centre/stories/2025/"],
              ["T07 News Listing by Year", "EDS", "Simple",
               "Single dynamic news-list block bound to a year index.",
               "https://www.westernsydney.edu.au/news-centre/stories/2026"],
              ["T08 Expert Opinion / Student Spotlight", "EDS", "Simple-Medium",
               "Article variant of the news story template. ~134 pages.",
               "https://www.westernsydney.edu.au/news-centre/expert-opinion/"],
              ["T09 Events Listing / Archive", "EDS", "Medium",
               "Table-based (and index-driven) event lists per microsite. ~884 pages.",
               "https://www.westernsydney.edu.au/babylab/events/previous-events"],
              ["T10 People / Team Listing", "EDS", "Simple-Medium",
               "Grid/cards of people and leadership. ~250+ pages.",
               "https://www.westernsydney.edu.au/challenging-racism-project/our-people"],
              ["T11 Nav & Footer Fragments", "EDS", "Medium",
               "Navigation authored per-microsite as a page (header-fragment meta); shared /footer fragment. Not a visitor page type but a reusable authored template.",
               "https://www.westernsydney.edu.au/ics/nav\nhttps://www.westernsydney.edu.au/footer"],
              ["T12 AEM Adaptive Form", "AEM", "Complex",
               "fd/af runtime clientlibs; multi-step wizard + accordion forms. No native EDS equivalent.",
               "(embedded on homepage / enquiry pages)"],
          ],
          widths=[3.6, 1.2, 1.8, 6.6, 5.2])
doc.add_page_break()

# ---------- 3. Blocks ----------
h(doc, "3. Blocks / Components Catalog", 1)
body(doc, "Reusable blocks, grouping visual variants of one content model as a single block with variants "
          "noted. The EDS blocks below are already live in production; the AEM Core Components are the "
          "residual homepage/corporate set to be reproduced as EDS blocks.")
h(doc, "3.1 Edge Delivery Blocks (interior estate — already live)", 2)
add_table(doc,
          ["Block", "Complexity", "Behaviour & Functionality", "Example URL"],
          [
              ["Hero (variants: tint-low, title-secondary, white title/description)", "Medium",
               "Full-width banner image + heading + intro + CTA.", "https://www.westernsydney.edu.au/schools/grs"],
              ["Carousel", "Medium",
               "Rotating slides: image link + heading + sub-heading + 'Find out more' link.", "https://www.westernsydney.edu.au/ics"],
              ["Cards (variants: light-grey, off-white, crimson; btn-outline)", "Medium",
               "Responsive grid of image + heading + text + optional button. Colour variants are the main design lever.", "https://www.westernsydney.edu.au/schools/grs"],
              ["Columns (incl. button-list variant)", "Simple-Medium",
               "Multi-column layout; button-list renders a row of CTAs; section background via section-metadata.", "https://www.westernsydney.edu.au/schools/grs"],
              ["Table (variants: dynamic-width, no-header, no-border, two-column-table)", "Simple-Medium",
               "Flexible table used for data AND as a curated media-list layout (image % + text %). Heavily reused for story lists.", "https://www.westernsydney.edu.au/news-centre"],
              ["News-list (dynamic)", "Complex",
               "Index/query-driven feed: config rows for query, limit, pagination; renders latest stories automatically.", "https://www.westernsydney.edu.au/news-centre"],
              ["Accordions", "Medium",
               "Expand/collapse panels, each title + rich content (images, paragraphs). Used within news stories.", "https://www.westernsydney.edu.au/news-centre/stories/2025/"],
              ["Global Navigation (fragment block)", "Medium",
               "Nested mega-menu authored as a page + section-metadata config (search link, sign-in label, red CTA, logo). Per-microsite.", "https://www.westernsydney.edu.au/ics/nav"],
              ["Global Footer (fragment block)", "Simple-Medium",
               "Shared columns layout + social/contact icons (facebook, instagram, linkedin, tiktok, phone, location).", "https://www.westernsydney.edu.au/footer"],
              ["Default content (headings, paragraphs, links, picture, buttons)", "Simple",
               "Standard EDS auto-decorated content; images auto-optimised (webp/jpg srcset via the aem.live pipeline).", "https://www.westernsydney.edu.au/ics/about-us"],
              ["Section metadata (styling primitive)", "Simple",
               "Sets section background/name and carries block config (not a visual block).", "https://www.westernsydney.edu.au/schools/grs"],
              ["Video embed (YouTube)", "Simple",
               "Embedded YouTube player.", "https://www.westernsydney.edu.au/"],
          ],
          widths=[4.4, 1.8, 7.6, 4.6])
h(doc, "3.2 Classic AEM Core Components (homepage / corporate — to be rebuilt as EDS blocks)", 2)
add_table(doc,
          ["Component", "Complexity", "Behaviour & Functionality", "Example URL"],
          [
              ["News & Events (cmp-news-events)", "Complex",
               "Bespoke composite: news + events cards, image gallery, and a calendar. Closest EDS analog = dynamic news-list + table.", "https://www.westernsydney.edu.au/"],
              ["Promo Band (cmp-promo-block--band)", "Medium", "Full-width promotional band (multiple instances).", "https://www.westernsydney.edu.au/"],
              ["Button (cmp-button ×6 variants)", "Simple", "solid / outline / text / black / white / red variants.", "https://www.westernsydney.edu.au/"],
              ["Container / Title / Text / Image", "Simple", "Standard AEM Core WCM building blocks.", "https://www.westernsydney.edu.au/"],
              ["Experience Fragment (corporate nav & footer)", "Medium",
               "Classic AEM XF for the corporate shell — to be re-expressed as EDS nav/footer fragments (already done for microsites).", "https://www.westernsydney.edu.au/"],
              ["Algolia Search (custom web components)", "Complex",
               "algolia-autocomplete / -configure / -hit-template reading 5 indices incl. courses.", "https://www.westernsydney.edu.au/future"],
          ],
          widths=[4.4, 1.8, 7.6, 4.6])
doc.add_page_break()

# illustrative screenshots
h(doc, "3.3 Illustrative Screenshots", 2)
shot(doc, "home", "Figure 1: Corporate Homepage (T01) — classic AEM, cmp-news-events, promo bands.")
shot(doc, "institute", "Figure 2: Research-institute microsite landing /ics (T03) — EDS hero, carousel, cards.")
doc.add_page_break()
shot(doc, "news-story", "Figure 3: News story detail (T06) — EDS rich-text + accordions.")
shot(doc, "news-centre", "Figure 4: News Centre hub (T05) — EDS dynamic news-list + curated tables.")
doc.add_page_break()

# ---------- 4. Page counts ----------
h(doc, "4. Page Counts by Template", 1)
body(doc, "Counts derive from URL-pattern analysis of the 9,925-page sitemap. Because the interior estate is "
          "already on Edge Delivery, 'migration type' below reflects the residual effort: RE-PLATFORM for the "
          "classic-AEM corporate pages, and VALIDATE/RE-HOST for pages already on EDS (confirm block fidelity, "
          "metadata, redirects — largely automatable).")
add_table(doc,
          ["Template", "Stack", "Est. Pages", "Migration Type", "Notes / Rationale"],
          [
              ["T06 News Story Detail", "EDS", "~1,870", "Automatic (validate)", "Uniform rich-text + images; already EDS. Bulk-verify."],
              ["T09 Events Listing / Archive", "EDS", "~884", "Automatic (validate)", "Table/index-driven lists; already EDS."],
              ["T10 People / Team Listing", "EDS", "~250+", "Automatic (validate)", "Card/grid people lists; already EDS."],
              ["T08 Expert Opinion / Spotlight", "EDS", "~134", "Automatic (validate)", "Article variant; already EDS."],
              ["T05 News Centre Hub", "EDS", "~15", "Semi-Automatic", "Dynamic news-list config to verify."],
              ["T07 News Listing by Year", "EDS", "~7", "Automatic (validate)", "Single dynamic block per year."],
              ["T03 Microsite Landing", "EDS", "~300-400", "Semi-Automatic", "40+ centres; hero/carousel/cards + nav fragment per site; verify config."],
              ["T04 Generic Content / Rich-Text", "EDS", "~6,300", "Automatic (validate)", "Long tail of about/policy/info/publication pages; already EDS default content."],
              ["T11 Nav & Footer Fragments", "EDS", "~40+", "Semi-Automatic", "One nav fragment per microsite + shared footer."],
              ["T01/T02 Corporate Homepage & Marketing", "AEM", "~10-40", "Manual (re-platform)", "Classic AEM; bespoke widgets, promo bands, Algolia. Not in sitemap."],
              ["T12 AEM Adaptive Forms", "AEM", "~10-30", "Manual (re-build)", "Adaptive Forms have no native EDS equivalent."],
              ["TOTAL (sitemap, EDS)", "EDS", "9,925", "-", "Overwhelmingly already on Edge Delivery."],
          ],
          widths=[4.2, 1.1, 2.0, 2.6, 6.4])
h(doc, "4.1 Migration Automatability Summary", 2)
add_table(doc,
          ["Category", "Est. Pages", "Includes"],
          [
              ["Already on EDS — validate/re-host (largely automatic)", "~9,800", "News, events, people, microsite content, generic pages. Confirm block fidelity, metadata, redirects."],
              ["Semi-automatic (config verification)", "~350-450", "Microsite landings, news-centre hub, nav/footer fragments."],
              ["Manual re-platform / re-build", "~20-70", "Classic-AEM corporate homepage & marketing pages, Adaptive Forms, Algolia search block."],
          ],
          widths=[6.4, 2.2, 8.0])
body(doc, "Because ~99% of sitemap pages already render on Edge Delivery, the dominant activity is validation "
          "and cut-over rather than content re-authoring. The concentrated manual effort is the small but "
          "complex corporate/marketing shell, forms, and search.", size=9, italic=True, color=GREY)
doc.add_page_break()

# ---------- 5. Integrations ----------
h(doc, "5. Integrations Analysis", 1)
body(doc, "Integrations confirmed from rendered DOM, inline scripts, and the Adobe DTM library. Actual IDs "
          "captured where visible. Enterprise systems on separate subdomains are noted as out-of-scope links.")
add_table(doc,
          ["Integration", "Type", "Complexity", "Evidence (ID / domain)", "Where"],
          [
              ["Adobe DTM / Launch", "Tag manager", "Medium",
               "assets.adobedtm.com/…satelliteLib…; injected by EDS delayed.js", "All pages"],
              ["Adobe Audience Manager", "Analytics / DMP", "Medium", "demdex; org 864008A3573B47607F000101@AdobeOrg", "via DTM"],
              ["Google Tag Manager", "Tag manager", "Simple", "GTM-KDSCL5, GTM-NCPWGGZR", "via DTM"],
              ["Google Ads / Floodlight", "Ad pixel", "Simple", "AW-11286271048, AW-937284738, AW-18181713940; dc-432…", "via DTM"],
              ["Microsoft Clarity", "Analytics", "Simple", "clarity.ms / scripts.clarity.ms", "Homepage"],
              ["Meta / Facebook Pixel", "Ad pixel", "Simple", "fbevents.js; 5 pixel IDs incl. custom CoursePageView event", "via DTM"],
              ["TikTok Pixel", "Ad pixel", "Simple", "analytics.tiktok.com; 4 pixel IDs", "via DTM"],
              ["Snapchat Pixel", "Ad pixel", "Simple", "snaptr init dfe8aba0-…", "via DTM"],
              ["LinkedIn Insight", "Ad pixel", "Simple", "snap.licdn.com; partner_id 9227514", "via DTM"],
              ["Everest / Conversant", "Ad pixel", "Simple", "everestjs.net last-event-tag", "via DTM"],
              ["Spotify Pixel", "Ad pixel", "Simple", "pixel.byspotify.com", "Homepage"],
              ["OneTrust", "Consent (CMP)", "Complex", "cdn-au.onetrust.com/consent/6cfd50d2-…/otSDKStub.js", "via DTM"],
              ["Algolia", "Search", "Complex",
               "app NY25AYQTZ8; indices wsu_prod_courses/articles/content_library/main_pages; custom web components", "Homepage, /future"],
              ["Genesys PureCloud", "Live chat", "Complex", "apps.mypurecloud.com.au/genesys-bootstrap (region au)", "via DTM"],
              ["Qualtrics", "Survey / intercept", "Complex", "…qualtrics.com/WRSiteInterceptEngine; ZN_* zone IDs", "via DTM"],
              ["AEM Adaptive Forms", "Forms", "Complex", "core/fd/af-clientlibs runtime; wizard + accordion; dompurify", "Classic AEM pages"],
              ["YouTube", "Media embed", "Simple", "youtube.com/embed/…", "Homepage"],
              ["AEM EDS RUM", "Telemetry", "Native", "rum.hlx.page; sampleRUM in scripts.js", "All pages"],
              ["jQuery 3.6.0", "Library", "Simple (retire)", "ajax.googleapis.com/…jquery-3.6.0", "Homepage, /future"],
              ["Slick Carousel 1.8.1", "Library", "Simple (rebuild)", "cdnjs…/slick-carousel/1.8.1", "Homepage"],
              ["Cloud.typography (Hoefler)", "Web fonts", "Simple", "cloud.typography.com", "/future"],
              ["Okta (SSO)", "Auth", "Complex (out of scope)", "login.westernsydney.edu.au (Okta OIDC/MFA)", "login subdomain"],
              ["Handbook system", "Course/handbook", "Complex (out of scope)", "handbook. / hbook. / studenthandbook.westernsydney.edu.au", "Linked"],
              ["Applications portal", "Application", "Complex (out of scope)", "apply.westernsydney.edu.au", "Linked"],
              ["Library (Springshare LibCal)", "Library", "Complex (out of scope)", "library.westernsydney.edu.au → api3-au.libcal.com", "library subdomain"],
              ["ServiceNow", "ITSM", "Medium (out of scope)", "wsu.service-now.com/it", "Footer"],
              ["GoGlobal / Careers / Online portals", "Portals", "Medium (out of scope)", "goglobal. / careers. / online.westernsydney.edu.au", "Linked"],
          ],
          widths=[3.9, 2.3, 2.2, 5.6, 3.2], font=8, hfont=8.5)
h(doc, "5.1 Migration Notes", 2)
body(doc, "Trivial to carry across (already proven on EDS): the entire tag/pixel layer — WSU's EDS delayed.js "
          "already injects the same Adobe DTM library, so GTM, Audience Manager and all ad pixels transfer for "
          "free. YouTube → standard embed block. jQuery, Slick and cloud.typography should be retired/rebuilt "
          "natively.", size=10)
body(doc, "Needs real re-engineering: Algolia search (rebuild custom web components as an EDS search/"
          "autocomplete block against app NY25AYQTZ8, repoint indexing at EDS content); AEM Adaptive Forms "
          "(no EDS equivalent — rebuild or iframe-embed); Genesys chat & Qualtrics intercepts (load via "
          "delayed.js but re-verify page-targeting against the URL structure); OneTrust (re-validate category "
          "blocking). Enterprise subdomains (handbook, apply, library, login/Okta, goglobal) are out of CMS "
          "scope — preserve links and SSO hand-offs.", size=10)
doc.add_page_break()

# ---------- 6. Complex use cases ----------
h(doc, "6. Complex Use Cases & Observations", 1)
body(doc, "Behaviours and edge cases needing special attention during migration, quantified. [C] = confirmed "
          "against live pages/sitemap; [E] = estimated.")
add_table(doc,
          ["Use Case", "Instances", "Where Found", "Why It's Complex"],
          [
              ["Dual-stack site (AEM + EDS in parallel)", "Whole estate [C]",
               "Homepage/marketing = AEM (cms_type=AEM); ~9,900 interior pages = EDS (cms_type=EDS)",
               "Two rendering stacks, two component vocabularies, and a live cut-over boundary. Cross-stack nav/consistency and redirect management are the central risk."],
              ["Bespoke news-events composite component", "1 (homepage) [C]",
               "cmp-news-events (news + events + calendar + gallery)",
               "Single custom AEM component fuses four content types + a calendar. No 1:1 EDS block; must be rebuilt from news-list + table + gallery."],
              ["Dynamic news-list (index-driven feeds)", "News-centre + microsites [C]",
               "/news-centre and year listings; config rows for query/limit/pagination",
               "Query/pagination logic bound to content indexes; feeds must be reproduced and re-pointed to EDS-published content, and kept in sync."],
              ["Algolia federated search (5 indices)", "Site-wide search + course search [C]",
               "app NY25AYQTZ8; wsu_prod_courses/articles/content_library/main_pages",
               "Custom web components and a course index power study discovery. Rebuilding as an EDS block plus re-indexing EDS content is the largest single integration lift."],
              ["AEM Adaptive Forms", "~10-30 forms [E]",
               "fd/af runtime; wizard + accordion; homepage enquiry forms",
               "Multi-step conditional forms with validation/submit actions; no native EDS equivalent — each needs rebuild or embed, with backend endpoints re-tested."],
              ["Per-microsite navigation fragments", "40+ [C]",
               "/{microsite}/nav pages + section-metadata (search link, sign-in, red CTA)",
               "Each of 40+ centres/schools has its own authored mega-menu fragment; migration must preserve per-site IA, not a single global nav."],
              ["Research-institute microsite sprawl", "40+ microsites, ~5,000+ pages [C]",
               "/ics, /marcs, /thri, /nicmhri, /hie, /babylab, … each with about-us/people/research/projects/publications/events/news",
               "Large volume of near-identical sub-structures with local variation; template detection and QA at scale dominate effort."],
              ["Large ad-pixel / recruitment martech farm", "15+ trackers [C]",
               "Meta ×5, TikTok ×4, Snapchat, LinkedIn, Google Ads ×3, Floodlight, Spotify; custom CoursePageView",
               "Consent-gated conversion tracking tied to student recruitment; category/zone targeting must be re-validated post cut-over under OneTrust."],
              ["Genesys PureCloud live chat + Qualtrics intercepts", "Site-wide [C]",
               "apps.mypurecloud.com.au; Qualtrics ZN_* zones",
               "Session/stateful widgets with page-targeting rules held in DTM; behaviour must be re-verified against new URLs."],
              ["Enterprise systems on separate subdomains", "6+ systems [C]",
               "handbook, apply, library→LibCal, login→Okta, goglobal, careers, service-now",
               "Out of CMS scope but deeply linked; broken-link/SSO-handoff risk if inbound links and deep links aren't inventoried and preserved."],
              ["Bot protection (HTTP 406 to non-browsers)", "Whole site [C]",
               "406 Not Acceptable without browser headers",
               "Automated crawling/import tooling must send browser headers; affects any migration/QA automation and third-party indexing."],
          ],
          widths=[3.4, 2.2, 4.4, 5.6], font=8, hfont=8.5)
doc.add_page_break()

# ---------- Appendix ----------
h(doc, "Appendix: Template Screenshots", 1)
body(doc, "Representative page per template family (captured August 2026, 1440px viewport).", size=9, italic=True, color=GREY)
figs = [
    ("news-list", "News listing by year (T07)"),
    ("expert-op", "Expert opinion article (T08)"),
    ("inst-about", "Microsite content / about page (T04)"),
    ("school", "School landing (T03) — hero, cards, button-list"),
    ("events-list", "Events archive listing (T09)"),
    ("people", "People / team listing (T10)"),
    ("research", "Research/projects listing (T03/T04)"),
    ("content-deep", "Deep content page — inherent requirements (T04)"),
]
i = 5
for key, label in figs:
    shot(doc, key, f"Figure {i}: {label}.")
    doc.add_paragraph()
    i += 1

note = doc.add_paragraph()
r = note.add_run("Prepared from live inspection of representative pages across both platforms (August 2026). "
                 "Page counts sourced from the site sitemap. This assessment is a planning document; a fixed "
                 "migration plan requires a template/block workshop and confirmation of the AEM-vs-EDS cut-over roadmap with WSU.")
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

out = os.path.join(BASE, "WesternSydney_Migration_Analysis.docx")
doc.save(out)
print("Saved:", out)
print("Size: %.1f KB" % (os.path.getsize(out) / 1024))
